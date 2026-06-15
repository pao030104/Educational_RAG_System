"""
Word 文档加载器模块
=====================
基于 python-docx 的 Word (.docx) 文档加载器，提取文本段落、表格和图片文字。

工作流程:
    ┌──────────────┐    ┌──────────────────┐    ┌──────────────────┐
    │ 打开 .docx    │───→│ 按块迭代          │───→│ 分类处理:        │
    │              │    │ (iter_block_items)│    │ • 段落 → 直接取  │
    │              │    │                  │    │   含图片OCR      │
    └──────────────┘    └──────────────────┘    │ • 表格 → 遍历提取 │
                                                └──────────────────┘

块迭代器 (iter_block_items):
    按照 Word 文档 XML 结构中的元素顺序迭代段落和表格。
    这是 python-docx 官方推荐的方式，能保持内容的原始阅读顺序。

    Word 文档内部结构 (XML):
        <w:body>              ← Document.body
            <w:p>              ← CT_P → Paragraph (段落)
                <w:r>          ← Run (文本运行)
                <w:drawing>    ← 内嵌图片
            </w:p>
            <w:tbl>            ← CT_Tbl → Table (表格)
                <w:tr>         ← Row (行)
                <w:tc>         ← Cell (单元格)
                    <w:p>      ← 单元格内的段落
        </w:body>

图片 OCR:
    段落中嵌入的图片通过 XPath 查询定位:
        - './/pic:pic': 查找所有图片元素
        - './/a:blip/@r:embed': 获取图片的关系ID
        - part.related_parts[img_id]: 通过关系ID获取图片二进制数据
"""

from typing import Iterator                         # 类型注解
from edu_ocr import get_ocr                         # OCR 引擎
from tqdm import tqdm                               # 进度条

# ---- python-docx 组件 ----
from docx.table import _Cell, Table                 # 表格和单元格
from docx.oxml.table import CT_Tbl                  # 表格 XML 元素
from docx.oxml.text.paragraph import CT_P            # 段落 XML 元素
from docx.text.paragraph import Paragraph            # 段落对象
from docx import Document as Docu1                   # 文档读取（别名避开冲突）
from docx.document import Document as Docu2          # 文档类型检查用
from docx import ImagePart                           # 图片部件

from PIL import Image                                # 图片处理
from io import BytesIO                               # 字节流转图片
import numpy as np                                   # 数组转换

from langchain_core.documents import Document         # LangChain Document
from langchain_core.document_loaders import BaseLoader  # 基类


class OCRDOCLoader(BaseLoader):
    """
    Word (.docx) 文档加载器，提取段落、表格和图片中的所有文字。

    遵循 LangChain BaseLoader 接口。

    属性:
        filepath (str): .docx 文件路径。

    使用示例:
        >>> loader = OCRDOCLoader(filepath="/path/to/document.docx")
        >>> docs = loader.load()
        >>> text = docs[0].page_content
    """

    def __init__(self, filepath: str) -> None:
        """
        初始化 Word 加载器。

        参数:
            filepath (str): .docx 文件的路径。
        """
        self.filepath = filepath

    def lazy_load(self) -> Iterator[Document]:
        """
        延迟加载模式：提取文本后返回单个 Document。

        Returns:
            Iterator[Document]
        """
        line = self.doc2text(self.filepath)
        yield Document(page_content=line, metadata={"source": self.filepath})

    def doc2text(self, filepath):
        """
        从 .docx 文件提取全部文本内容。

        处理流程:
            1. 打开文档
            2. 按 XML 元素顺序迭代段落和表格
            3. 段落: 提取直接文本 + 嵌入图片 OCR
            4. 表格: 遍历行 → 单元格 → 段落

        参数:
            filepath (str): .docx 文件路径。

        返回:
            str: 提取的全部文本内容。
        """
        ocr = get_ocr()  # 获取 OCR 引擎

        # 打开 Word 文档
        doc = Docu1(filepath)
        resp = ""  # 累积全部文本

        def iter_block_items(parent):
            """
            按文档 XML 元素顺序迭代段落和表格。

            这是 python-docx 官方推荐的内容按序迭代方案。
            直接使用 doc.paragraphs + doc.tables 无法保证顺序，
            因为 Word 文档中段落和表格是交错排列的。

            参数:
                parent: Document 或 _Cell 对象。

            Yields:
                Paragraph 或 Table 对象，按文档中的出现顺序。
            """
            # 根据 parent 类型获取 XML body 元素
            if isinstance(parent, Docu2):
                # Document 对象 → 获取 <w:body>
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                # 单元格对象 → 获取 <w:tc>
                parent_elm = parent._tc
            else:
                raise ValueError("OCRDOCLoader parse fail")

            # 按 XML 子元素顺序迭代
            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    # CT_P (w:p) → 段落
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    # CT_Tbl (w:tbl) → 表格
                    yield Table(child, parent)

        # ---- 创建进度条 ----
        b_unit = tqdm(
            total=len(doc.paragraphs) + len(doc.tables),
            desc="OCRDOCLoader block index: 0"
        )

        # ---- 按序处理文档块 ----
        for i, block in enumerate(iter_block_items(doc)):
            b_unit.set_description(
                "OCRDOCLoader  block index: {}".format(i)
            )
            b_unit.refresh()

            if isinstance(block, Paragraph):
                # ---- 处理段落 ----
                resp += block.text.strip() + "\n"

                # ---- 提取段落中的嵌入图片 ----
                # 使用 XPath 查找段落元素中的所有图片
                # './/pic:pic': 查找所有命名空间为 pic 的 pic 元素
                images = block._element.xpath('.//pic:pic')
                for image in images:
                    # 获取图片的关系 ID (r:embed 属性)
                    for img_id in image.xpath('.//a:blip/@r:embed'):
                        # 通过关系 ID 获取图片部件
                        part = doc.part.related_parts[img_id]
                        if isinstance(part, ImagePart):
                            # 打开图片二进制数据
                            image = Image.open(BytesIO(part._blob))
                            # OCR 识别
                            result, _ = ocr(np.array(image))
                            if result:
                                ocr_result = [line[1] for line in result]
                                resp += "\n".join(ocr_result)

            elif isinstance(block, Table):
                # ---- 处理表格 ----
                for row in block.rows:            # 遍历行
                    for cell in row.cells:         # 遍历单元格
                        for paragraph in cell.paragraphs:  # 遍历单元格段落
                            resp += paragraph.text.strip() + "\n"

            b_unit.update(1)

        return resp


if __name__ == '__main__':
    # 测试 Word 加载器
    # 请将路径修改为你本地的 Word 文件路径
    import sys
    test_file = sys.argv[1] if len(sys.argv) > 1 else './samples/ocr_02.docx'
    docx_loader = OCRDOCLoader(filepath=test_file)
    doc = docx_loader.load()
    print(doc)
